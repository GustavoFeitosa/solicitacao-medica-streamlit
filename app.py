import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io

# ===========================
# LISTAS DE EXAMES
# ===========================
exames_lab = [
    "Hemograma Completo",
    "Sódio, Potássio, Uréia, Creatinina, Ácido úrico",
    "15 OH Vitamina D3",
    "Cálcio iônico",
    "Colesterol total e frações",
    "Glicemia, Hb-glicada",
    "AST, ALT, CPK",
    "Ferritina",
    "NT-próBNP",
    "TSH, T4-livre",
    "Sumário de Urina",
    "Relação albumina/creatinina em amostra isolada de urina",
    "Lp(a)",
    "VHS",
    "PCR de alta sensibilidade"
]

exames_imagem = [
    "Teste Ergométrico",
    "RX de tórax em PA e Perfil",
    "Ultrassom de Abdome Total",
    "Ecocardiograma Bidimensional com Doppler Colorido",
    "ECG de repouso",
    "Holter de 24 horas",
    "MAPA de 24 horas",
    "Cintilografia do miocárdio sob repouso e estresse físico",
    "Cintilografia do miocárdio sob repouso e estresse farmacológico",
    "AngioTC de coronárias",
    "TC de tórax para escore de cálcio coronariano",
    "Doppler arterial de membros inferiores",
    "Doppler venoso de membros inferiores",
    "US com Doppler de carótidas e vertebrais",
    "US de tireoide",
    "Endoscopia Digestiva Alta",
    "Colonoscopia"
]

# ===========================
# FUNÇÃO PARA CRIAR RECEITUÁRIO
# ===========================
def criar_receituario(paciente, cid, justificativa, lista_exames, titulo):

    doc = Document()

    # ---------- Cabeçalho ----------
    h = doc.add_paragraph("CONSULTÓRIO CARDIOLÓGICO")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(14)

    n = doc.add_paragraph("Dr. Gustavo Feitosa – Cardiologista")
    n.alignment = WD_ALIGN_PARAGRAPH.CENTER
    n.runs[0].font.size = Pt(12)

    doc.add_paragraph("")

    # ---------- Título ----------
    t = doc.add_paragraph(titulo)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(14)

    doc.add_paragraph("")

    # ---------- Paciente ----------
    p = doc.add_paragraph(f"Para Sr(a). {paciente}")
    p.runs[0].font.size = Pt(12)

    # ---------- Exames ----------
    doc.add_paragraph("Solicito:\n").runs[0].font.size = Pt(12)

    for ex in lista_exames:
        linha = doc.add_paragraph(f"• {ex}")
        linha.runs[0].font.size = Pt(12)

    # Justificativa – duas linhas acima do CID
    if justificativa.strip():
        doc.add_paragraph("\n")
        j = doc.add_paragraph(f"Justificativa: {justificativa}")
        j.runs[0].font.size = Pt(12)

    # CID – duas linhas acima da data
    doc.add_paragraph("\n")
    cid_par = doc.add_paragraph(f"CID 10: {cid}")
    cid_par.runs[0].font.size = Pt(12)

    # ---------- Data ----------
    data = datetime.now().strftime("%d/%m/%Y")
    d = doc.add_paragraph(f"\nSalvador/BA, {data}")
    d.runs[0].font.size = Pt(12)

    # ---------- Assinatura ----------
    assinatura = doc.add_paragraph("\n_______________________________")
    assinatura.runs[0].font.size = Pt(12)

    info = doc.add_paragraph(
        "Dr. Gustavo Feitosa – Cardiologista\nCRM/BA 21730 – RQE 21919"
    )
    info.runs[0].font.size = Pt(11)

    doc.add_paragraph("")

    # ---------- Rodapé ----------
    rod1 = doc.add_paragraph(
        "Centro Médico Aliança, sala 211: Av. Juracy Magalhães Júnior, 2096 –\n"
        "Rio Vermelho, Salvador - BA, 41920-180\n"
        "Tel: 71 21084686"
    )
    rod1.runs[0].font.size = Pt(10)

    # Logo HA
    ha_paragraph = doc.add_paragraph()
    ha_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ha_paragraph.add_run().add_picture("assets/logo_ha.png", width=Pt(65))

    doc.add_paragraph("")  # Espaço

    rod2 = doc.add_paragraph(
        "Centro Médico Cárdio Pulmonar, sala 501: Rua Ponciano Oliveira, 157 – Rio Vermelho,\n"
        "Salvador – BA, 41920-275\n"
        "Tel: 71 30344598 / 71 30344599"
    )
    rod2.runs[0].font.size = Pt(10)

    # Logo HCP (selo dourado + azul juntos)
    hcp_paragraph = doc.add_paragraph()
    hcp_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hcp_paragraph.add_run().add_picture("assets/logo_hcp.png", width=Pt(110))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ===========================
# INTERFACE STREAMLIT
# ===========================

st.image("assets/logo_zumtec.png", width=260)
st.markdown("### 📄 Gerador de Solicitações Médicas – Dr. Gustavo Feitosa")
st.markdown("---")

paciente = st.text_input("Nome completo do paciente")
cid = st.text_input("CID (ex: I-10, I-25.1)")
justificativa = st.text_area("Justificativa (opcional)")

# ===========================
# LABORATORIAIS – MARCAR TODOS
# ===========================
st.markdown("### 🧪 Exames Laboratoriais")

select_all_lab = st.checkbox("Selecionar todos os exames laboratoriais")

cols_lab = st.columns(2)
lab_selecionados = []

for i, exame in enumerate(exames_lab):
    col = cols_lab[i % 2]
    marcado = col.checkbox(exame, value=select_all_lab)
    if marcado:
        lab_selecionados.append(exame)

# Campo livre extra
extra_lab = st.text_area("Exames laboratoriais adicionais (digite um por linha)")
if extra_lab.strip():
    extras = [x.strip() for x in extra_lab.split("\n") if x.strip()]
    lab_selecionados.extend(extras)

# ===========================
# IMAGEM – COM CAMPO LIVRE
# ===========================
st.markdown("### 🩻 Exames de Imagem / Complementares")

cols_img = st.columns(2)
img_selecionados = []

for i, exame in enumerate(exames_imagem):
    col = cols_img[i % 2]
    if col.checkbox(exame):
        img_selecionados.append(exame)

extra_img = st.text_area("Exames de imagem/complementares adicionais (um por linha)")
if extra_img.strip():
    extras = [x.strip() for x in extra_img.split("\n") if x.strip()]
    img_selecionados.extend(extras)

# ===========================
# BOTÃO
# ===========================
if st.button("Gerar Solicitações"):
    if paciente.strip() == "" or cid.strip() == "":
        st.error("Preencha nome e CID.")
    else:

        # ---- Laboratoriais ----
        if lab_selecionados:
            doc_lab = criar_receituario(
                paciente, cid, justificativa, lab_selecionados,
                "SOLICITAÇÃO DE EXAMES LABORATORIAIS"
            )
            st.download_button(
                "📥 Baixar Solicitação Laboratorial",
                data=doc_lab,
                file_name="solicitacao_laboratorial.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # ---- Exames de imagem – 1 arquivo por exame ----
        for exame in img_selecionados:
            doc_img = criar_receituario(
                paciente, cid, justificativa, [exame],
                "SOLICITAÇÃO DE EXAME COMPLEMENTAR"
            )
            filename = exame.replace(" ", "_").lower() + ".docx"
            st.download_button(
                f"📥 Solicitação – {exame}",
                data=doc_img,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.success("Solicitações geradas com sucesso!")
